#!/usr/bin/env python3
"""Convert ZATCA markdown docs to professional Word documents with RTL Arabic support."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def add_styled_paragraph(doc, text, size=11, bold=False, color=None, align=None, rtl=True, space_after=6):
    p = doc.add_paragraph()
    if rtl:
        set_rtl(p)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Arial'
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._r.get_or_add_rPr()
    rtl_el = OxmlElement('w:rtl')
    rPr.append(rtl_el)
    return p


def add_heading(doc, text, level=1, color='2C5559'):
    sizes = {1: 22, 2: 18, 3: 14, 4: 12}
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._r.get_or_add_rPr()
    rtl_el = OxmlElement('w:rtl')
    rPr.append(rtl_el)


def parse_inline(text):
    """Strip markdown inline formatting (bold, code, links) into plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = re.sub(r'^[-*]\s+', '• ', text)
    text = re.sub(r'^\d+\.\s+', '', text)
    return text.strip()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        set_cell_bg(cell, '2C5559')
        p = cell.paragraphs[0]
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(parse_inline(h))
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Arial'
        rPr = run._r.get_or_add_rPr()
        rPr.append(OxmlElement('w:rtl'))
    # Body rows
    n_cols = len(headers)
    for r_idx, row in enumerate(rows):
        # Pad or truncate to match header count
        row = (row + [''] * n_cols)[:n_cols]
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            set_rtl(p)
            run = p.add_run(parse_inline(val))
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            rPr = run._r.get_or_add_rPr()
            rPr.append(OxmlElement('w:rtl'))


def md_to_docx(md_path, docx_path, title):
    doc = Document()
    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        # RTL section
        sectPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        sectPr.append(bidi)

    # Cover/title
    add_heading(doc, title, level=1, color='2C5559')

    text = Path(md_path).read_text(encoding='utf-8')
    lines = text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip horizontal rules
        if re.match(r'^-{3,}$', line) or re.match(r'^={3,}$', line):
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            add_heading(doc, parse_inline(m.group(2)), level=level)
            i += 1
            continue

        # Tables
        if line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1]):
            headers = [c.strip() for c in line.strip('|').split('|')]
            i += 2  # skip header + separator
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                row = [c.strip() for c in lines[i].strip('|').split('|')]
                rows.append(row)
                i += 1
            add_table(doc, headers, rows)
            continue

        # Code blocks (fenced)
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string('444444')
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (could be list item)
        text_line = parse_inline(line)
        is_list = line.lstrip().startswith(('- ', '* ', '• '))
        is_numbered = re.match(r'^\s*\d+\.\s+', line)
        size = 11
        if is_list or is_numbered:
            add_styled_paragraph(doc, text_line, size=size, space_after=3)
        else:
            add_styled_paragraph(doc, text_line, size=size)
        i += 1

    doc.save(docx_path)
    print(f"✓ Created: {docx_path}")


if __name__ == '__main__':
    docs_dir = Path(__file__).parent
    md_to_docx(
        docs_dir / 'ZATCA_SDD.md',
        docs_dir / 'ZATCA_SDD.docx',
        'وثيقة الوصف الفني للنظام — محسوب (Mahsoob SDD)'
    )
    md_to_docx(
        docs_dir / 'ZATCA_VISIT_CHECKLIST.md',
        docs_dir / 'ZATCA_VISIT_CHECKLIST.docx',
        'قائمة التحضير لزيارة هيئة الزكاة والضريبة والجمارك'
    )
    print("\n✅ Done! Files saved in:", docs_dir)
